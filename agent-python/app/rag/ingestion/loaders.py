"""Document loaders — PDF, Word, Markdown, HTML, Text, CSV and S3/MinIO."""

from __future__ import annotations

import csv
import io
import logging
import os
import tempfile

from app.rag.ingestion.types import Document

logger = logging.getLogger(__name__)


class TextLoader:
    """Load plain text and CSV files."""

    SUPPORTED_EXTENSIONS = {".txt", ".csv", ".json", ".log", ".xml", ".yaml", ".yml"}

    def load(self, path: str) -> list[Document]:
        ext = os.path.splitext(path)[1].lower()
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        if ext == ".csv":
            content = self._format_csv(content)

        doc = Document(
            content=content,
            metadata={
                "source": os.path.basename(path),
                "format": ext.lstrip("."),
                "path": path,
            },
        )
        return [doc]

    def _format_csv(self, content: str) -> str:
        try:
            reader = csv.reader(io.StringIO(content))
            rows = list(reader)
            if not rows:
                return content
            header = rows[0]
            lines = [",".join(header)]
            for row in rows[1:]:
                lines.append(" | ".join(f"{h}: {v}" for h, v in zip(header, row)))
            return "\n".join(lines)
        except Exception:
            return content


class PDFLoader:
    """Load PDF files via PyPDF2."""

    SUPPORTED_EXTENSIONS = {".pdf"}

    def load(self, path: str) -> list[Document]:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF loading. Install: pip install PyPDF2")

        reader = PdfReader(path)
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Page {i + 1}]\n{text.strip()}")

        content = "\n\n".join(pages)
        doc = Document(
            content=content,
            metadata={
                "source": os.path.basename(path),
                "format": "pdf",
                "page_count": len(reader.pages),
                "path": path,
            },
        )
        return [doc]


class DocxLoader:
    """Load Word (.docx) files."""

    SUPPORTED_EXTENSIONS = {".docx"}

    def load(self, path: str) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required. Install: pip install python-docx")

        docx = DocxDocument(path)
        paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        doc = Document(
            content=content,
            metadata={
                "source": os.path.basename(path),
                "format": "docx",
                "paragraph_count": len(paragraphs),
                "path": path,
            },
        )
        return [doc]


class MarkdownLoader:
    """Load Markdown files with optional YAML frontmatter extraction."""

    SUPPORTED_EXTENSIONS = {".md", ".markdown", ".mdx"}

    def load(self, path: str) -> list[Document]:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()

        metadata: dict = {"source": os.path.basename(path), "format": "md", "path": path}
        content = raw

        if raw.startswith("---"):
            parts = raw.split("---", 2)
            if len(parts) >= 3:
                frontmatter = parts[1].strip()
                content = parts[2].strip()
                metadata.update(self._parse_frontmatter(frontmatter))

        doc = Document(content=content, metadata=metadata)
        return [doc]

    def _parse_frontmatter(self, text: str) -> dict[str, str]:
        result: dict[str, str] = {}
        for line in text.split("\n"):
            line = line.strip()
            if ":" in line:
                key, _, val = line.partition(":")
                result[key.strip()] = val.strip().strip("\"'")
        return result


class HTMLLoader:
    """Load HTML files, extracting text via BeautifulSoup."""

    SUPPORTED_EXTENSIONS = {".html", ".htm"}

    def load(self, path: str) -> list[Document]:
        with open(path, "rb") as f:
            raw = f.read()
        return self.load_bytes(raw, source=os.path.basename(path))

    def load_bytes(self, data: bytes, source: str = "inline.html") -> list[Document]:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 is required. Install: pip install beautifulsoup4")

        soup = BeautifulSoup(data, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)

        doc = Document(
            content=text,
            metadata={"source": source, "format": "html"},
        )
        return [doc]


class S3Loader:
    """Load documents from S3/MinIO, delegating to the appropriate format loader."""

    def __init__(self, endpoint: str | None = None, access_key: str | None = None, secret_key: str | None = None):
        self._endpoint = endpoint or os.getenv("S3_ENDPOINT", os.getenv("MINIO_ENDPOINT", ""))
        self._access_key = access_key or os.getenv("S3_ACCESS_KEY", os.getenv("MINIO_ACCESS_KEY", ""))
        self._secret_key = secret_key or os.getenv("S3_SECRET_KEY", os.getenv("MINIO_SECRET_KEY", ""))
        self._bucket = os.getenv("S3_BUCKET", "enterprise-ai-docs")

    def load(self, s3_uri: str) -> list[Document]:
        try:
            from minio import Minio
        except ImportError:
            raise ImportError("minio is required for S3 loading. Install: pip install minio")

        bucket, key = self._parse_uri(s3_uri)
        client = Minio(
            self._endpoint.replace("https://", "").replace("http://", ""),
            access_key=self._access_key,
            secret_key=self._secret_key,
            secure=self._endpoint.startswith("https://"),
        )

        with tempfile.NamedTemporaryFile(suffix=os.path.splitext(key)[1], delete=False) as tmp:
            client.fget_object(bucket, key, tmp.name)
            tmp_path = tmp.name

        try:
            fmt_loader = loader_for(key)
            docs = fmt_loader.load(tmp_path)
            for doc in docs:
                doc.metadata["s3_uri"] = s3_uri
                doc.metadata["s3_bucket"] = bucket
            return docs
        finally:
            os.unlink(tmp_path)

    def _parse_uri(self, uri: str) -> tuple[str, str]:
        prefix = "s3://"
        if uri.startswith(prefix):
            uri = uri[len(prefix):]
        parts = uri.split("/", 1)
        if len(parts) != 2:
            raise ValueError(f"Invalid S3 URI: {uri}")
        return parts[0], parts[1]


def loader_for(path: str) -> TextLoader | PDFLoader | DocxLoader | MarkdownLoader | HTMLLoader:
    """Return the appropriate loader for a file based on its extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext in PDFLoader.SUPPORTED_EXTENSIONS:
        return PDFLoader()
    if ext in DocxLoader.SUPPORTED_EXTENSIONS:
        return DocxLoader()
    if ext in MarkdownLoader.SUPPORTED_EXTENSIONS:
        return MarkdownLoader()
    if ext in HTMLLoader.SUPPORTED_EXTENSIONS:
        return HTMLLoader()
    return TextLoader()
